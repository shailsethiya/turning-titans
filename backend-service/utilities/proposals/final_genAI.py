
import docx
from langchain.embeddings.sentence_transformer import SentenceTransformerEmbeddings
from langchain.text_splitter import CharacterTextSplitter
from langchain.vectorstores import Chroma
from langchain.chains import RetrievalQA
from flask import current_app as app
import openai, os
from langchain.llms import AzureOpenAI

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

def generate_proposal(filepath_rfp,filepath_output):
    # read rfp from filepath
    text_rfp = getText(filepath_rfp)
    # generate proposal
    ## fetch knowledgebase
    embedding_function = SentenceTransformerEmbeddings(model_name="all-MiniLM-L6-v2")
    chroma_path_offering = os.path.join(os.path.dirname(app.config['UPLOAD_FOLDER']),"offerings","refract" ,"chroma_db")
    print(chroma_path_offering)
    db = Chroma(persist_directory=chroma_path_offering, embedding_function=embedding_function)

    openai.api_type = "azure"
    openai.api_key = "840a080abe584b299d700c777616e216"
    openai.api_base = "https://refract-openai.openai.azure.com/"
    openai.api_version = "2023-03-15-preview"

    dep_name="refract-gpt-35-turbo" # 'Refract-ada001'
    model_name="refract-gpt-35-turbo" # 'Refract-ada001'

    os.environ["OPENAI_API_TYPE"] = "azure"
    os.environ["OPENAI_API_VERSION"] = "2023-03-15-preview"
    os.environ["OPENAI_API_BASE"] = "https://refract-openai.openai.azure.com/"
    os.environ["DEPLOYMENT_NAME"] = dep_name
    os.environ["OPENAI_API_KEY"] = "840a080abe584b299d700c777616e216"

    llm = AzureOpenAI(deployment_name=dep_name, model_name=model_name, temperature=0,max_tokens=3000)

    prompt = text_rfp + """\nAbove is the content from an RFP, for which I have to create a proposal. Your task is to create sections of my proposal. As response return a python dictionary with key as
    section heading and value as a prompt. The prompt will be used to generate the content for the respective section and should be with respect to a product called Refract. Refract is a self-served, unified data analytics platform that abstracts the operationalization of the AI lifecycle, behind the scene
    Make sure the promp is elaborate enough to generate some content for the respective section"""

    response = llm(prompt)

    import ast
    # Find the index of the first '{' and the last '}' in the string
    start_index = response.find('{')
    end_index = response.rfind('}') + 1

    # Extract the dictionary string from the input string
    dictionary_string = response[start_index:end_index]

    # Parse the dictionary string using ast.literal_eval() to get the dictionary
    try:
        extracted_dictionary = ast.literal_eval(dictionary_string)
        print("Extracted Dictionary:", extracted_dictionary)
    except SyntaxError:
        extracted_dictionary = {'Introduction': 'Introduce Refract and the company submitting the proposal. Explain briefly what Refract is and how it can help the bank.', 'Understanding of the problem statement': 'Explain the current challenges faced by the bank in building and maintaining their analytics infrastructure. Highlight the limitations of their current infrastructure and how it is affecting their business.', 'Proposed Solution': "Explain how Refract can help the bank overcome the challenges they are facing. Highlight the key features of Refract that make it a good fit for the bank's requirements.", 'Technical Approach': 'Explain the technical approach that will be taken to build the MAP infrastructure using Refract. Highlight the key components of the infrastructure and how they will be integrated with Refract.', 'Project Plan': 'Provide a detailed project plan that outlines the key milestones and deliverables for the project. Explain how the project will be managed and how progress will be tracked.', 'Team Structure': 'Provide details of the team that will be working on the project. Explain the roles and responsibilities of each team member and how they will work together to deliver the project.', 'Cost Estimation': 'Provide a detailed cost estimation for the project. Break down the costs into modular and granular components so that it is possible to consider finalizing a contract on selected components. Explain how the costs have been arrived at and how they compare to industry standards.', 'Conclusion': 'Summarize the proposal and reiterate the key benefits of using Refract to build the MAP infrastructure. Highlight the competitive advantage that the bank will gain by using Refract and how it will help them achieve their strategic objectives.'}
        print("Invalid dictionary string in the input.")
        print(response)

    

    qa_chain = RetrievalQA.from_chain_type(
        llm=llm,
        chain_type="stuff",
        retriever=db.as_retriever(),
        return_source_documents=True
    )
    proposal = {}
    for section, prompt in extracted_dictionary.items():
        res = qa_chain({"query":prompt})["result"]
        print(res)
        proposal[section] = res
    
    print(proposal)
    #save as docx file at "filepath_output + /proposal.docx" 
    from docx import Document

    document = Document()

    document.add_heading('Refract: Empowering Enterprise AI/ML Adoption', 0)
    for section, content in proposal.items():
        document.add_heading(section, level=1)
        document.add_paragraph(content)
    document.add_page_break()

    document.save(os.path.join(filepath_output, 'proposal.docx'))
    return None 
